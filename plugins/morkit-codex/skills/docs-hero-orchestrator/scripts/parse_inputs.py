"""Parse input documents (PDF/Excel/Docx/OpenSpec specs) into a raw bundle.

This script does **structural** extraction only — text, tables, paragraphs.
Semantic classification (which chunk is an FR vs NFR vs Screen) is delegated
to Claude in the orchestrator (LLM-assisted, see phase-02 §2.4).

Output JSON:
    {
      "meta": {...},
      "raw_content": {
        "pdf_files": [...],
        "excel_files": [...],
        "docx_files": [...],
        "openspec_specs": [...]
      }
    }

CLI:
    parse-inputs.py --inputs-dir inputs/ [--openspec-dir openspec/]
                    --language EN --project-name X --output raw.json

Public API:
    parse_inputs(inputs_dir, openspec_dir, language, project_name) -> dict
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document  # python-docx

sys.path.insert(0, str(Path(__file__).resolve().parent))

from lib.pdf_extractor import extract_pdf  # noqa: E402

PDF_EXT = {".pdf"}
EXCEL_EXT = {".xlsx", ".xlsm"}
DOCX_EXT = {".docx"}


def parse_inputs(
    inputs_dir: str | Path | None,
    openspec_dir: str | Path | None,
    language: str = "EN",
    project_name: str = "Untitled Project",
) -> dict[str, Any]:
    """Walk inputs_dir + openspec_dir, extract structured content."""
    raw_content: dict[str, list[dict[str, Any]]] = {
        "pdf_files": [],
        "excel_files": [],
        "docx_files": [],
        "openspec_specs": [],
    }
    warnings: list[str] = []

    if inputs_dir:
        inputs_path = Path(inputs_dir)
        if not inputs_path.exists():
            warnings.append(f"inputs-dir not found: {inputs_path}")
        else:
            for file in sorted(inputs_path.rglob("*")):
                if not file.is_file():
                    continue
                ext = file.suffix.lower()
                try:
                    if ext in PDF_EXT:
                        raw_content["pdf_files"].append(_extract_pdf(file))
                    elif ext in EXCEL_EXT:
                        raw_content["excel_files"].append(_extract_excel(file))
                    elif ext in DOCX_EXT:
                        raw_content["docx_files"].append(_extract_docx(file))
                except Exception as exc:
                    warnings.append(f"failed {file}: {exc}")

    if openspec_dir:
        specs_path = Path(openspec_dir) / "specs"
        if specs_path.exists():
            for spec_file in sorted(specs_path.glob("*.md")):
                raw_content["openspec_specs"].append(
                    {
                        "path": str(spec_file),
                        "content": spec_file.read_text(encoding="utf-8"),
                    }
                )
        else:
            warnings.append(f"openspec/specs not found: {specs_path}")

    return {
        "meta": {
            "project_name": project_name,
            "version": "1.0",
            "language": language,
        },
        "raw_content": raw_content,
        "warnings": warnings,
    }


def _extract_pdf(path: Path) -> dict[str, Any]:
    content = extract_pdf(path)
    return {
        "path": str(path),
        "page_count": content.page_count,
        "text": content.text,
        "tables": content.tables,
        "warnings": content.extraction_warnings,
    }


def _extract_excel(path: Path) -> dict[str, Any]:
    wb = openpyxl.load_workbook(filename=str(path), data_only=True, read_only=True)
    sheets: list[dict[str, Any]] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            if all(cell is None for cell in row):
                continue
            rows.append(["" if cell is None else str(cell) for cell in row])
        if not rows:
            continue
        headers, data_rows = rows[0], rows[1:]
        sheets.append(
            {
                "name": sheet_name,
                "headers": headers,
                "rows": data_rows,
                "row_count": len(data_rows),
            }
        )
    wb.close()
    return {"path": str(path), "sheets": sheets}


def _extract_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [
        {"style": p.style.name if p.style else "Normal", "text": p.text}
        for p in doc.paragraphs
        if p.text.strip()
    ]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        table_rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(table_rows)
    return {"path": str(path), "paragraphs": paragraphs, "tables": tables}


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    p.add_argument("--inputs-dir", help="Directory with PDF/Excel/Docx files")
    p.add_argument("--openspec-dir", help="OpenSpec root (with specs/ inside)")
    p.add_argument("--language", default="EN", choices=["JP", "EN", "VN"])
    p.add_argument("--project-name", default="Untitled Project")
    p.add_argument("--output", required=True)
    args = p.parse_args()

    if not args.inputs_dir and not args.openspec_dir:
        p.error("Provide at least one of --inputs-dir or --openspec-dir")

    result = parse_inputs(
        inputs_dir=args.inputs_dir,
        openspec_dir=args.openspec_dir,
        language=args.language,
        project_name=args.project_name,
    )
    Path(args.output).write_text(
        json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    rc = result["raw_content"]
    print(
        f"Parsed: {len(rc['pdf_files'])} PDF, {len(rc['excel_files'])} Excel, "
        f"{len(rc['docx_files'])} Docx, {len(rc['openspec_specs'])} OpenSpec specs"
        f" -> {args.output}",
        file=sys.stderr,
    )
    if result["warnings"]:
        print(f"Warnings: {len(result['warnings'])}", file=sys.stderr)
        for w in result["warnings"]:
            print(f"  - {w}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
