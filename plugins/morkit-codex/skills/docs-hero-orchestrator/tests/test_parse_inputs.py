"""Tests for parse_inputs.py."""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import openpyxl
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from parse_inputs import parse_inputs  # noqa: E402


def _create_sample_excel(path: Path) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "DataItems"
    ws.append(["ID", "Field", "Type", "Length"])
    ws.append(["DATA-001", "email", "VARCHAR", 255])
    ws.append(["DATA-002", "name", "VARCHAR", 100])
    ws2 = wb.create_sheet("FRs")
    ws2.append(["ID", "Name", "Description"])
    ws2.append(["FR-001", "Login", "User login flow"])
    wb.save(str(path))


def _create_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Requirements", level=1)
    doc.add_paragraph("Project requirements for the system.")
    doc.add_heading("FR-001 Login", level=2)
    doc.add_paragraph("User can sign in with email and password.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Type"
    table.rows[1].cells[0].text = "email"
    table.rows[1].cells[1].text = "string"
    doc.save(str(path))


def test_parse_excel_only():
    with tempfile.TemporaryDirectory() as td:
        inputs = Path(td) / "inputs"
        inputs.mkdir()
        _create_sample_excel(inputs / "data.xlsx")

        result = parse_inputs(
            inputs_dir=inputs, openspec_dir=None, language="EN", project_name="Test"
        )

    rc = result["raw_content"]
    assert len(rc["excel_files"]) == 1
    excel = rc["excel_files"][0]
    assert len(excel["sheets"]) == 2
    assert excel["sheets"][0]["name"] == "DataItems"
    assert excel["sheets"][0]["headers"] == ["ID", "Field", "Type", "Length"]
    assert excel["sheets"][0]["rows"][0] == ["DATA-001", "email", "VARCHAR", "255"]


def test_parse_docx_only():
    with tempfile.TemporaryDirectory() as td:
        inputs = Path(td) / "inputs"
        inputs.mkdir()
        _create_sample_docx(inputs / "spec.docx")

        result = parse_inputs(
            inputs_dir=inputs, openspec_dir=None, language="EN", project_name="Test"
        )

    rc = result["raw_content"]
    assert len(rc["docx_files"]) == 1
    docx = rc["docx_files"][0]
    paragraph_texts = [p["text"] for p in docx["paragraphs"]]
    assert any("FR-001" in t for t in paragraph_texts)
    assert len(docx["tables"]) == 1
    assert docx["tables"][0][0] == ["Field", "Type"]


def test_parse_openspec_specs_only():
    with tempfile.TemporaryDirectory() as td:
        openspec = Path(td) / "openspec"
        (openspec / "specs").mkdir(parents=True)
        (openspec / "specs" / "auth.md").write_text("# Auth\n\nLogin spec.")

        result = parse_inputs(
            inputs_dir=None, openspec_dir=openspec, language="EN", project_name="Test"
        )

    rc = result["raw_content"]
    assert len(rc["openspec_specs"]) == 1
    assert "Login spec" in rc["openspec_specs"][0]["content"]


def test_parse_combined_inputs_and_openspec():
    with tempfile.TemporaryDirectory() as td:
        inputs = Path(td) / "inputs"
        inputs.mkdir()
        _create_sample_excel(inputs / "data.xlsx")

        openspec = Path(td) / "openspec"
        (openspec / "specs").mkdir(parents=True)
        (openspec / "specs" / "auth.md").write_text("# Auth")

        result = parse_inputs(
            inputs_dir=inputs, openspec_dir=openspec, language="JP", project_name="P"
        )

    assert result["meta"]["language"] == "JP"
    assert result["meta"]["project_name"] == "P"
    assert len(result["raw_content"]["excel_files"]) == 1
    assert len(result["raw_content"]["openspec_specs"]) == 1


def test_missing_input_dir_warns():
    with tempfile.TemporaryDirectory() as td:
        nonexistent = Path(td) / "does-not-exist"
        result = parse_inputs(
            inputs_dir=nonexistent, openspec_dir=None, language="EN", project_name="P"
        )
    assert any("not found" in w for w in result["warnings"])
    assert result["raw_content"]["pdf_files"] == []


def test_meta_present_in_output():
    with tempfile.TemporaryDirectory() as td:
        inputs = Path(td) / "inputs"
        inputs.mkdir()
        result = parse_inputs(
            inputs_dir=inputs, openspec_dir=None, language="VN", project_name="MyProj"
        )
    assert result["meta"]["project_name"] == "MyProj"
    assert result["meta"]["language"] == "VN"
    assert result["meta"]["version"] == "1.0"
